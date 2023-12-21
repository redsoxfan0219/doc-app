from docx import Document
import pandas as pd

## This function needs to be a method for the 
## Document / DocxTemplate objects

## temporarily creating document object

def create_table_from_source(document: Document,
                            source_path: str, 
                            worksheet=None,
                            header=True) -> Document:
        
    source_dataframe = __create_df_based_on_file_type(source_path=source_path)
    number_rows = source_dataframe.shape[0]
    number_cols = source_dataframe.shape[1]
    
    table = document.add_table(rows=1, cols=number_cols)
    table.style = "Table Grid"
    
    # Add header row to the table and set it to bold
    header_cells = table.rows[0].cells
    for i, col_name in enumerate(source_dataframe.columns.tolist()):
        run = header_cells[i].paragraphs[0].add_run(col_name)
        run.bold = True
        
    # Fill rows with data
    for index, row in source_dataframe.iterrows():
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            row_cells[idx].text = str(value)
    
    # Set bold for header row only
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    document.save('demo.docx')
                
    # create table
    
    # Get the number of rows and columns from the Pandas object
    
    
    # Set the header row if a header has been provided; don't set header if not
    # Use the create table functions in 
        
def __create_df_based_on_file_type(source_path:str,
                                   sep = ",") -> pd.DataFrame:
    
    if source_path.endswith(".xlsx"):
        df = pd.read_excel(source_path, worksheet=worksheet)
        return df
    elif source_path.endswith(".csv"):
        df = pd.read_csv(source_path, sep=sep, header=0)
        return df
    elif source_path.endswith(".parquet"):
        df = pd.read_parquet()
        return df
    elif source_path.endswith(".json"):
        df = pd.read_json()
        return df
    else:
        print("Invalid file type.")
        return None
    

if __name__ == "__main__":
    document = Document()
    create_table_from_source(document, "dummy_data.csv")
    
    
    
    
    
        
    
    
    
        
        
    
    
